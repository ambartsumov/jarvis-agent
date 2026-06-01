"""
PDS-Ultimate File Tools
========================
Инструменты для работы с файлами: Excel, Word, PDF, OCR.

ФУНКЦИИ:
- Создание Excel
- Чтение Excel
- Создание Word
- Чтение PDF
- Создание отчётов
- OCR изображений

ARCHITECTURE:
- openpyxl для Excel
- python-docx для Word
- PyPDF2 для PDF
- easyocr для OCR
- Async-safe operations
"""

from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from pds_ultimate.config import config, logger
from pds_ultimate.core.tools import Tool, ToolParameter, ToolResult, ToolRegistry

# ─── File Tools ─────────────────────────────────────────────────────────────


async def tool_create_excel(
    filename: str,
    data: str,
    sheet_name: str = "Sheet1",
) -> ToolResult:
    """
    Создать Excel файл из данных.
    
    Формат данных: JSON или CSV в строке.
    """
    try:
        from openpyxl import Workbook
        
        import json
        
        # Parse data
        try:
            rows = json.loads(data)
            if not isinstance(rows, list):
                rows = [rows]
        except json.JSONDecodeError:
            # Try CSV parsing
            lines = data.strip().split("\n")
            headers = lines[0].split(",")
            rows = []
            for line in lines[1:]:
                values = line.split(",")
                rows.append(dict(zip(headers, values)))
        
        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name
        
        # Write headers
        if rows:
            headers = list(rows[0].keys())
            for col, header in enumerate(headers, 1):
                ws.cell(row=1, column=col, value=header)
            
            # Write data
            for row_idx, row_data in enumerate(rows, 2):
                for col_idx, header in enumerate(headers, 1):
                    ws.cell(row=row_idx, column=col_idx, value=row_data.get(header))
        
        # Save
        os.makedirs("data/exports", exist_ok=True)
        filepath = f"data/exports/{filename}"
        if not filepath.endswith(".xlsx"):
            filepath += ".xlsx"
        
        wb.save(filepath)
        
        return ToolResult(
            "create_excel",
            True,
            f"✅ Excel создан: {filepath} ({len(rows)} строк)",
            data={"filepath": filepath, "rows": len(rows)},
        )
        
    except ImportError:
        return ToolResult(
            "create_excel",
            False,
            "openpyxl не установлен (pip install openpyxl)",
            error="openpyxl not installed",
        )
    except Exception as e:
        logger.error(f"tool_create_excel failed: {e}")
        return ToolResult("create_excel", False, "", error=str(e))


async def tool_read_excel(
    filepath: str,
    limit: int = 100,
) -> ToolResult:
    """
    Прочитать Excel файл.
    """
    try:
        from openpyxl import load_workbook
        
        if not os.path.exists(filepath):
            return ToolResult(
                "read_excel",
                False,
                f"Файл не найден: {filepath}",
            )
        
        wb = load_workbook(filepath, read_only=True)
        ws = wb.active
        
        # Read headers
        headers = [cell.value for cell in ws[1]]
        
        # Read data
        rows = []
        for row_idx, row in enumerate(ws.iter_rows(min_row=2, max_row=limit + 1), 1):
            row_data = {}
            for col_idx, cell in enumerate(row):
                if col_idx < len(headers):
                    row_data[headers[col_idx]] = cell.value
            rows.append(row_data)
        
        wb.close()
        
        # Format output
        output = f"📊 {filepath} ({len(rows)} строк):\n\n"
        for i, row in enumerate(rows[:10], 1):
            output += f"{i}. {row}\n"
        if len(rows) > 10:
            output += f"... и ещё {len(rows) - 10} строк"
        
        return ToolResult(
            "read_excel",
            True,
            output,
            data={"filepath": filepath, "rows": rows, "count": len(rows)},
        )
        
    except ImportError:
        return ToolResult(
            "read_excel",
            False,
            "openpyxl не установлен",
            error="openpyxl not installed",
        )
    except Exception as e:
        logger.error(f"tool_read_excel failed: {e}")
        return ToolResult("read_excel", False, "", error=str(e))


async def tool_create_word(
    filename: str,
    content: str,
    title: Optional[str] = None,
) -> ToolResult:
    """
    Создать Word документ.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        
        # Add title
        if title:
            heading = doc.add_heading(title, 0)
            heading.runs[0].font.size = Pt(18)
        
        # Add content
        for paragraph in content.split("\n\n"):
            doc.add_paragraph(paragraph)
        
        # Save
        os.makedirs("data/exports", exist_ok=True)
        filepath = f"data/exports/{filename}"
        if not filepath.endswith(".docx"):
            filepath += ".docx"
        
        doc.save(filepath)
        
        return ToolResult(
            "create_word",
            True,
            f"✅ Word создан: {filepath}",
            data={"filepath": filepath},
        )
        
    except ImportError:
        return ToolResult(
            "create_word",
            False,
            "python-docx не установлен (pip install python-docx)",
            error="python-docx not installed",
        )
    except Exception as e:
        logger.error(f"tool_create_word failed: {e}")
        return ToolResult("create_word", False, "", error=str(e))


async def tool_read_pdf(
    filepath: str,
    pages: Optional[str] = None,
) -> ToolResult:
    """
    Прочитать PDF файл.
    """
    try:
        import PyPDF2
        
        if not os.path.exists(filepath):
            return ToolResult(
                "read_pdf",
                False,
                f"Файл не найден: {filepath}",
            )
        
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            
            # Parse pages
            if pages:
                page_nums = [int(p) - 1 for p in pages.split(",")]
            else:
                page_nums = range(len(reader.pages))
            
            text = ""
            for page_num in page_nums:
                if 0 <= page_num < len(reader.pages):
                    text += reader.pages[page_num].extract_text()
        
        output = f"📄 {filepath} ({len(text)} символов):\n\n{text[:2000]}"
        if len(text) > 2000:
            output += "\n... (обрезано)"
        
        return ToolResult(
            "read_pdf",
            True,
            output,
            data={"filepath": filepath, "text": text[:5000]},
        )
        
    except ImportError:
        return ToolResult(
            "read_pdf",
            False,
            "PyPDF2 не установлен (pip install PyPDF2)",
            error="PyPDF2 not installed",
        )
    except Exception as e:
        logger.error(f"tool_read_pdf failed: {e}")
        return ToolResult("read_pdf", False, "", error=str(e))


async def tool_create_report(
    report_type: str,
    data: str,
    filename: Optional[str] = None,
) -> ToolResult:
    """
    Создать отчёт (Excel/PDF/Word).
    
    Типы: finance, orders, inventory, custom
    """
    try:
        from pds_ultimate.core.llm_engine import llm_engine
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{report_type}_report_{timestamp}"
        
        # Generate report content with LLM
        prompt = f"Создай профессиональный отчёт типа '{report_type}' на основе данных:\n{data}"
        
        content = await llm_engine.chat(
            message=prompt,
            task_type="generate_report",
            temperature=0.3,
        )
        
        # Determine format
        if report_type in ("finance", "orders", "inventory"):
            # Excel
            return await tool_create_excel(filename, content)
        else:
            # Word
            return await tool_create_word(filename, content)
        
    except Exception as e:
        logger.error(f"tool_create_report failed: {e}")
        return ToolResult("create_report", False, "", error=str(e))


async def tool_ocr_image(
    image_path: str,
    language: str = "ru",
) -> ToolResult:
    """
    Распознать текст на изображении (OCR).
    
    Поддерживает: чеки, накладные, документы, скриншоты.
    """
    try:
        import easyocr
        
        if not os.path.exists(image_path):
            return ToolResult(
                "ocr_image",
                False,
                f"Файл не найден: {image_path}",
            )
        
        # Initialize reader
        reader = easyocr.Reader([language], gpu=False)
        
        # Recognize text
        results = reader.readtext(image_path)
        
        text = "\n".join([r[1] for r in results])
        confidence = sum(r[2] for r in results) / max(1, len(results))
        
        output = f"📷 OCR: {image_path}\n\n{text}"
        output += f"\n\nУверенность: {confidence:.1%}"
        
        return ToolResult(
            "ocr_image",
            True,
            output,
            data={"text": text, "confidence": confidence, "filepath": image_path},
        )
        
    except ImportError:
        return ToolResult(
            "ocr_image",
            False,
            "easyocr не установлен (pip install easyocr)",
            error="easyocr not installed",
        )
    except Exception as e:
        logger.error(f"tool_ocr_image failed: {e}")
        return ToolResult("ocr_image", False, "", error=str(e))


# ─── Tool Registration ───────────────────────────────────────────────────────

def register_file_tools(registry: ToolRegistry) -> None:
    """Зарегистрировать file инструменты."""
    
    registry.register(
        Tool(
            name="create_excel",
            description="Создать Excel файл из данных (JSON/CSV)",
            parameters=[
                ToolParameter("filename", "string", "Имя файла"),
                ToolParameter("data", "string", "Данные (JSON или CSV)"),
                ToolParameter("sheet_name", "string", "Название листа", default="Sheet1", required=False),
            ],
            handler=tool_create_excel,
            category="file",
        )
    )
    
    registry.register(
        Tool(
            name="read_excel",
            description="Прочитать Excel файл",
            parameters=[
                ToolParameter("filepath", "string", "Путь к файлу"),
                ToolParameter("limit", "number", "Максимум строк", default=100, required=False),
            ],
            handler=tool_read_excel,
            category="file",
        )
    )
    
    registry.register(
        Tool(
            name="create_word",
            description="Создать Word документ",
            parameters=[
                ToolParameter("filename", "string", "Имя файла"),
                ToolParameter("content", "string", "Содержимое"),
                ToolParameter("title", "string", "Заголовок", required=False),
            ],
            handler=tool_create_word,
            category="file",
        )
    )
    
    registry.register(
        Tool(
            name="read_pdf",
            description="Прочитать PDF файл",
            parameters=[
                ToolParameter("filepath", "string", "Путь к файлу"),
                ToolParameter("pages", "string", "Номера страниц (через запятую)", required=False),
            ],
            handler=tool_read_pdf,
            category="file",
        )
    )
    
    registry.register(
        Tool(
            name="create_report",
            description="Создать отчёт (Excel/PDF/Word)",
            parameters=[
                ToolParameter("report_type", "string", "Тип отчёта (finance/orders/inventory/custom)"),
                ToolParameter("data", "string", "Данные для отчёта"),
                ToolParameter("filename", "string", "Имя файла", required=False),
            ],
            handler=tool_create_report,
            category="file",
        )
    )
    
    registry.register(
        Tool(
            name="ocr_image",
            description="Распознать текст на изображении (OCR)",
            parameters=[
                ToolParameter("image_path", "string", "Путь к изображению"),
                ToolParameter("language", "string", "Язык (ru/en)", default="ru", required=False),
            ],
            handler=tool_ocr_image,
            category="file",
        )
    )


__all__ = [
    "tool_create_excel",
    "tool_read_excel",
    "tool_create_word",
    "tool_read_pdf",
    "tool_create_report",
    "tool_ocr_image",
    "register_file_tools",
]
